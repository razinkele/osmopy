"""Generate the Baltic model calibration & status report (.docx), 2026-08-07."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor

OUT = Path("/home/razinka/osmopy/docs/baltic_model_status_report_2026-08-07.docx")

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)


def h(text, level):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    return p


def para(text, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    return p


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(hd)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(9.5)
    doc.add_paragraph()
    return t


def note(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.left_indent = Pt(18)
    return p


# ================= TITLE =================
t = doc.add_heading("Baltic OSMOSE Model — Calibration and Status Report", level=0)
for r in t.runs:
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
sub = doc.add_paragraph()
sr = sub.add_run(
    "Status as of 2026-08-07 · OSMOSE Python engine (osmopy), 9-species Baltic configuration"
)
sr.italic = True
para("Marine Research Institute, Klaipėda University")
note(
    "Companion documents: 'Pikeperch Overshoot in the Baltic OSMOSE Configuration' (2026-08-03, "
    "updated 2026-08-06) for the full percid investigation, and the LITERATURE folder for the "
    "monthly literature alerts. This report is the compact status recap."
)

# ================= EXEC SUMMARY =================
h("Executive summary", 1)
bullets(
    [
        "The current 9-species Baltic configuration certifies at 7 of 9 species inside their "
        "reference envelopes (50-year runs, five seeds, final-decade means). Under the new "
        "weight-aware verdict the headline is 5 of 5: every target backed by a real stock "
        "assessment is reproduced. The two failures — pikeperch (~58× over) and smelt (~5.7× "
        "over) — are the two lowest-confidence literature estimates in the reference file.",
        "No species collapses at equilibrium. Earlier 'collapse' verdicts (as low as 2/9) were an "
        "artifact of a defective persistence criterion that measured the seeding bootstrap "
        "transient instead of the equilibrium; the criterion was corrected and all prior "
        "certification notes audited (2026-08-01).",
        "The pikeperch overshoot investigation is concluded: the overshoot is structural and "
        "over-determined. Seven interventions across removal and production returned between "
        "−86% and +8% against a 58-fold gap, four intermediate explanations were withdrawn under "
        "scrutiny, and the biomass target itself was re-derived and retained — the gap is real, "
        "and this configuration cannot represent Baltic pikeperch at a defensible biomass on the "
        "current grid.",
        "Cross-engine fidelity with the Java reference engine improved substantially: a "
        "prey-axis defect in the diet diagnostic was found and repaired (it had been reporting "
        "every predator as 100% piscivorous), along with mortality-rate, header-parsing and "
        "diet-normalisation parity fixes. The test suite stands at ~4,390 tests.",
        "Literature watch (2026-07-28 alert): upstream OSMOSE has been static at v4.4.1 for "
        "three consecutive months — the 4.4.x parity port now has a stable target. The most "
        "actionable new paper is Voss & Quaas (2026) on temperature-dependent stock–recruitment "
        "for Western Baltic cod and herring. The ICES 2027 Baltic advice cycle is confirmed, so "
        "the reference envelopes for the pelagics have moved materially and need refreshing.",
    ]
)

# ================= CONFIGURATION =================
h("1. Configuration at a glance", 1)
table(
    ["Component", "Value"],
    [
        (
            "Focal species (9)",
            "cod_west, cod_east (disaggregated 2026-07-25), herring, sprat, flounder, "
            "perch, pikeperch, smelt, stickleback",
        ),
        ("Background predators", "grey seal, cormorant"),
        (
            "LTL resource groups (6)",
            "diatoms, dinoflagellates, micro-/meso-/macrozooplankton, benthos",
        ),
        ("Domain", "whole Baltic model domain (10–30°E, 54–66°N), 593 active cells"),
        ("Forcing", "CMEMS-derived LTL biomass and physics"),
        (
            "Engines",
            "pure-Python (NumPy/Numba, full Java parity: 14/14 EEC, 8/8 Bay of Biscay) "
            "and Java OSMOSE via subprocess",
        ),
        (
            "Certification protocol",
            "scripts/baltic_stability_certify.py — 50-year runs, seeds [42, 123, 7, 999, 2024], "
            "final-decade means vs data/baltic/reference/biomass_targets.csv",
        ),
    ],
)

# ================= CERT STATUS =================
h("2. Certification status (current configuration)", 1)
para(
    "Final-decade mean biomass against the reference envelopes. Weights are the confidence "
    "weights carried in the reference file (1.0 = category-1 ICES assessment, 0.5 = medium, "
    "0.2–0.3 = literature estimate poorly resolved at grid scale)."
)
table(
    ["Species", "Mean (t)", "Envelope (t)", "Weight", "Tier", "Status"],
    [
        ("cod_west", "~14,000", "4,000 – 25,000", "1.0", "assessed", "PASS"),
        ("cod_east", "83,122", "60,000 – 85,000", "1.0", "assessed", "PASS (2.3% headroom)"),
        ("herring", "2,600,112", "800,000 – 3,000,000", "1.0", "assessed", "PASS"),
        ("sprat", "1,059,941", "800,000 – 2,500,000", "1.0", "assessed", "PASS"),
        ("flounder", "~44,900", "20,000 – 100,000", "0.5", "assessed", "PASS"),
        ("perch", "45,382", "8,000 – 50,000", "0.2", "indicative", "PASS"),
        ("pikeperch", "1,453,313", "4,000 – 25,000", "0.2", "indicative", "FAIL (~58× over)"),
        ("smelt", "680,125", "20,000 – 120,000", "0.3", "indicative", "FAIL (~5.7× over)"),
        ("stickleback", "80,159", "50,000 – 500,000", "0.2", "indicative", "PASS"),
    ],
)
para(
    "Verdicts: assessed tier 5/5 PASS (the weight-aware headline) · all-species 7/9 (retained as "
    "the legacy figure for comparability with earlier notes) · indicative tier 2/4.",
    bold=True,
)
note(
    "Means are quoted from the five-seed certification where available and from single-seed "
    "diagnostic runs otherwise (cod_west, flounder are approximate mid-range values). "
    "cod_east is the tightest stock in the system at 2.3% headroom to its upper bound."
)
para(
    "The certifier became weight-aware on 2026-08-06: it now reads the confidence weight column "
    "from the reference file itself and reports assessed and indicative tiers separately, with "
    "the former all-species count retained as a labelled legacy figure. Eleven regression tests "
    "pin the behaviour. Previously all nine species were scored identically, which made the "
    "headline verdict substantially a statement about the two weakest targets rather than about "
    "the model."
)

# ================= CRITERION AUDIT =================
h("3. The persistence-criterion audit (2026-08-01)", 1)
para(
    "The certifier's 'persists' criterion used the whole-run biomass minimum, which on the "
    "Baltic is dominated by the deep transient dip of the seeding bootstrap, while the "
    "'in-envelope' criterion has always used the final-decade mean — the two halves of every "
    "verdict described different windows. Commit 556ba3d rescoped 'persists' to the final "
    "decade."
)
bullets(
    [
        "Every prior certification note was audited: rows reading 'persists ✗' with "
        "'in-envelope ✓' are bootstrap artifacts, not collapses. The same five stocks recur "
        "in every affected note: cod (both stocks after disaggregation), sprat, flounder, perch.",
        "Under the corrected criterion the current configuration moved from a reported 2/9 to "
        "7/9, and the 8-species pre-disaggregation baseline re-certified at 5/8 — in both cases "
        "the persist set now coincides with the in-envelope set, so no species collapses at "
        "equilibrium.",
        "An artifactual signal does not mean the decision it prompted was wrong. cod_east's "
        "natural mortality was lowered from M = 1.2546 to M = 0.9 partly on the faulty flag; the "
        "revert was tested and rejected — at the original M the stock does not collapse but "
        "settles ~5.6× below its ICES envelope (final-decade mean ~14.5–15.1 kt against "
        "60–85 kt). M = 0.9 stands on the mean, independent of the flag that prompted it.",
        "Decisions citing a persistence flag are being re-tested individually; configs behind "
        "older notes should be re-certified under the corrected criterion before their "
        "persistence verdicts are cited.",
    ]
)

# ================= CALIBRATION HISTORY =================
h("4. Calibration timeline (July–August 2026)", 1)
table(
    ["When", "Step", "Outcome"],
    [
        (
            "2026-07-21",
            "Baseline calibration (NSGA-II + GP surrogate, SALib sensitivity)",
            "8-species baseline, 5/8 in envelope; β-identifiability r² = 0.10–0.99",
        ),
        (
            "2026-07-25",
            "Cod east/west disaggregation (Phase 1)",
            "cod split into cod_west (sp0) + cod_east (sp8) with name-labelled 15×15 "
            "predation matrices; joint recalibration hit an apex-predation bottleneck and the "
            "split was retained as a flagged configuration change rather than a refit",
        ),
        (
            "2026-07-28",
            "Realistic percid removals",
            "percid fishing mortality (~44–46 kt total removals) plus cormorant predation "
            "close perch to its envelope; certified at five seeds with no regression",
        ),
        (
            "2026-07-28",
            "cod_east M correction",
            "M lowered 1.2546 → 0.9; revert tested 2026-08-01 and rejected (see Section 3)",
        ),
        (
            "2026-08-01",
            "Seeding-mode A/B (stock_recruitment vs linear) and criterion fix",
            "equilibrium means byte-identical across seeding modes — mode immaterial; "
            "persistence criterion corrected and all notes audited",
        ),
        (
            "2026-08-02/06",
            "Percid overshoot investigation and target re-derivation",
            "overshoot established as structural; targets retained; certifier made weight-aware",
        ),
    ],
)
para(
    "The recruitment-viability (RV) gate on cod_east remains load-bearing for its envelope PASS: "
    "a reference sweep put the valid tolerance band at 0.331–0.449 and the constraint is "
    "asymmetric, with ~2.2% headroom in the shipped profile. A surrogate-Bayesian UQ layer "
    "exists and is validated by self-consistency; the Baltic's literature-estimate targets are "
    "too weak to support ICES-target UQ directly."
)

# ================= PERCID =================
h("5. The percid overshoot: concluded as structural", 1)
para(
    "Pikeperch sits at ~1.45 Mt against a 4–25 kt target (~58×); smelt is ~5.7× over. Perch — "
    "same family, same coastal habitat, comparable fishery — is in envelope and served as the "
    "control throughout. The full investigation is in the companion report; the conclusions:"
)
bullets(
    [
        "Not an engine defect: consumption-to-biomass ratios at equilibrium are within "
        "published ranges for every species (pikeperch has the lowest Q/B in the system, 2.55); "
        "no conservation violation.",
        "Four explanations were reached and withdrawn under measurement or adversarial review: "
        "a resource subsidy (equilibrium diet is 36.9% resource, below perch), a spatial-supply "
        "mechanism (the corrected time-weighted herring overlap refutes it — perch has twice "
        "pikeperch's access and is in envelope), predation release (pikeperch is fourth-least "
        "predated; the three less-predated species are all in envelope), and a mis-scoped "
        "target (the reference file already declares whole-domain aggregate scope).",
        "The surviving mechanism — escape from every predator's size window before maturity "
        "(no predator can take a pikeperch above 36 cm; it matures at 40 cm) — is real but "
        "not sufficient: opening cod's window past maturity moved biomass +3.4%, the wrong way. "
        "Cod totals ~97 kt against pikeperch's 1.45 Mt; grey seal consumes ~58.5 kt/yr across "
        "all prey. No predator in the configuration has the biomass to regulate the stock.",
        "Seven interventions (accessibility cuts, F sweeps to 4.0, gear selectivity, size "
        "windows, reduced asymptotic length, recruitment ssbhalf cut 97.6×) returned −86% to "
        "+8% against the 58× gap. Biomass scales as roughly the 0.42 power of the recruitment "
        "parameter; reaching the target would require indefensible values.",
        "The target was re-derived and retained: ICES assesses no Baltic pikeperch, perch, "
        "smelt or stickleback stock, so these are literature estimates — but the perch target "
        "survives an independent production cross-check against the Curonian Lagoon anchor "
        "(Hansson et al., 2018), and pikeperch must sit below perch given its restricted, "
        "patchy distribution and the depleted state of both assessed Baltic stocks "
        "(Jakubavičiūtė et al., 2022). The 58× gap is real.",
    ]
)
para(
    "Reading: the configuration cannot represent Baltic pikeperch at a defensible biomass — "
    "which is what the reference file's own note about coarse-grid under-resolution "
    "anticipated. Representing percids as separate coastal stock units is the only change that "
    "would address this like-for-like and is out of scope on cost grounds; the weight-aware "
    "verdict keeps the failure visible but out of the assessed-tier headline. Smelt "
    "(second-most abundant fish of the eastern Gulf of Finland after herring; Sendek & "
    "Bogdanov, 2019) is a smaller instance of the same limitation."
)

# ================= ENGINE =================
h("6. Engine and cross-engine fidelity", 1)
bullets(
    [
        "Diet diagnostic defect repaired (#146): the prey axis of dietMatrix/predatorPressure "
        "was mislabelled (focal + background instead of schools + resources), silently dropping "
        "all resource columns — every predator reported 100% fish. Any diet-based conclusion "
        "predating commits e121c6d/3e320ff has been re-derived; a regression test now asserts "
        "no fish consumes less than its own body weight annually.",
        "Java cross-engine parity fixes: mortality-rate outputs were counts not rates (#140), "
        "output header parsing (#141), dietMatrix percentage normalisation (#144), and the "
        "predatorPressure time convention. The remaining sprat seeding gap (~7.1×) is isolated "
        "to the SeedingInterface lambda convention.",
        "Seeding mode (population.seeding.mode: stock_recruitment | linear) is configurable and "
        "measured immaterial at equilibrium (#143).",
        "Test suite: ~4,390 tests, green as of 2026-08-06, including 11 new weight-aware "
        "certifier tests.",
    ]
)

# ================= NEXT STEPS =================
h("7. Open items and next steps", 1)
table(
    ["Priority", "Item", "Rationale"],
    [
        (
            "High",
            "Ingest the ICES 2026 assessment cycle (advice for 2027) and refresh "
            "biomass_targets.csv",
            "cycle confirmed; reported directions: central Baltic herring +74% TAC, sprat "
            "+32%, both cod stocks zero catch — numbers are trade-press-sourced and must be "
            "re-read from the ICES advice sheets before use",
        ),
        (
            "High",
            "Temperature-dependent stock–recruitment in processes/reproduction.py",
            "Voss & Quaas (2026): ignoring temperature-dependent SR 'might trigger "
            "overexploitation'; osmopy already ingests CMEMS temperature but does not route it "
            "into reproduction; design jointly with the 4.4.x temperature-bioenergetics slice",
        ),
        (
            "High",
            "OSMOSE 4.4.x parity port (staged)",
            "upstream static at v4.4.1 for three consecutive months — the moving-target "
            "argument for deferral has expired",
        ),
        (
            "High",
            "WGSAM cod predation-mortality as a validation target",
            "2022 Baltic key run (1974–2021, ~64,000 cod stomachs) gives an independent check "
            "on emergent M2 on sprat and central Baltic herring",
        ),
        (
            "Medium",
            "Re-certify older configs under the corrected persistence criterion",
            "prior notes' persistence verdicts are unreliable until re-run (Section 3)",
        ),
        (
            "Medium",
            "Percid representation decision",
            "separate coastal stock units are the only like-for-like fix for the pikeperch/smelt "
            "limitation; currently excluded on cost grounds (cod E/W is the cautionary precedent)",
        ),
    ],
)

# ================= LITERATURE =================
h("8. Literature watch (from the LITERATURE folder)", 1)
para(
    "Monthly automated alerts run against scite, ICES/HELCOM, CRAN, ecopath.org and the "
    "osmose-model GitHub org; the latest is dated 2026-07-28 (window 2026-06-23 → 2026-07-28)."
)
bullets(
    [
        "Upstream OSMOSE: frozen for a third consecutive month at v4.4.1 (18 Jun 2026); "
        "no new org repos; homepage byte-identical. The whole v4.4.0 breaking-change set "
        "remains un-ported in osmopy (region-aware mortality, abundance-based fishing, "
        "stochastic maturity ogive, temperature bioenergetics, gradient movement, etc.).",
        "Most actionable paper: Voss & Quaas (2026), ICES JMS 83(4) — temperature-dependent "
        "stock–recruitment for Western Baltic cod and herring under RCP4.5/RCP8.5; herring "
        "recovers only after a ~4-year moratorium; cod catch potential stays below 5,000 t "
        "even under optimal management; MMEY dominates MMSY under climate change. (Note: the "
        "alert file credits Quaas as first author; the paper is Voss & Quaas.)",
        "ICES: 2027 Baltic catch advice confirmed published (see Section 7); the Baltic Sea "
        "Ecosystem Overview was read in full for the first time — late-1980s regime shift to "
        "sprat dominance, persistent hypoxia, warming and freshening are the qualitative "
        "behaviours the configuration should reproduce.",
        "Comparator models all unchanged: mizer 2.5.3, EwE 6.7 beta (build 18865), Atlantis "
        "with no public release. HELCOM HOLAS 4 (2022–2027) publishes in 2029.",
        "Carried comparators: Atlantis semi-automated calibration framework, global "
        "sensitivity-analysis of a complex marine ecosystem model, cross-ecosystem trophic "
        "transfer efficiency synthesis (all Ecological Modelling / Science Advances 2026).",
    ]
)

# ================= REFERENCES =================
h("References", 1)
para(
    "Hansson, S., Bergström, U., & Bonsdorff, E. (2018). Competition for the fish — fish "
    "extraction from the Baltic Sea by humans, aquatic mammals, and birds. ICES Journal of "
    "Marine Science, 75(3), 999–1008. https://doi.org/10.1093/icesjms/fsx207"
)
para(
    "Jakubavičiūtė, E., Arula, T., & Dainys, J. (2022). Status and future perspectives for "
    "pikeperch (Sander lucioperca) stocks in Europe. openRxiv preprint. "
    "https://doi.org/10.1101/2022.12.20.521162"
)
para(
    "Olsson, J., Jakubavičiūtė, E., Kaljuste, O., et al. (2019). The first large-scale "
    "assessment of three-spined stickleback (Gasterosteus aculeatus) biomass and spatial "
    "distribution in the Baltic Sea. ICES Journal of Marine Science, 76(6), 1653–1665. "
    "https://doi.org/10.1093/icesjms/fsz078"
)
para(
    "Sendek, D. S., & Bogdanov, D. V. (2019). European smelt Osmerus eperlanus in the eastern "
    "Gulf of Finland, Baltic Sea: Stock status and fishery. Journal of Fish Biology, 94(6), "
    "1001–1010. https://doi.org/10.1111/jfb.14009"
)
para(
    "Voss, R., & Quaas, M. F. (2026). Future fishing potential of cod and herring under "
    "climate change in the Western Baltic Sea. ICES Journal of Marine Science, 83(4), fsag033. "
    "https://doi.org/10.1093/icesjms/fsag033"
)

h("Source documents in the repository", 1)
bullets(
    [
        "docs/baltic_percid_overshoot_report_2026-08-03.docx — full percid investigation "
        "(updated 2026-08-06)",
        "docs/baltic_percid_target_rederivation_2026-08-06.md — target re-derivation",
        "docs/baltic_certification_reread_2026-08-01.md — persistence-criterion audit",
        "docs/baltic_8species_recert_corrected_criterion_2026-08-01.md — baseline re-certification",
        "docs/baltic_cod_east_M_revert_test_2026-08-01.md — M = 0.9 revert test",
        "docs/baltic_java_crossengine_fidelity_2026-07-30.md — cross-engine fidelity",
        "data/baltic/reference/biomass_targets.csv — targets, envelopes, confidence weights",
        "scripts/baltic_stability_certify.py — weight-aware certifier",
    ]
)

doc.save(OUT)
print(f"saved {OUT} ({OUT.stat().st_size} bytes)")
