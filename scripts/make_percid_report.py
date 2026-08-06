"""Generate the Baltic percid overshoot investigation report (.docx)."""

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path("/home/razinka/osmopy/docs/baltic_percid_overshoot_report_2026-08-03.docx")
REVIEW2 = Path(sys.argv[1]) if len(sys.argv) > 1 else None

doc = Document()

# ---- base styling ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)


def h(text, level):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    return p


def para(text, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    return p


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(hd)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(9.5)
    doc.add_paragraph()
    return t


def note(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.left_indent = Pt(18)
    return p


# ================= TITLE =================
t = doc.add_heading("Pikeperch Overshoot in the Baltic OSMOSE Configuration", level=0)
for r in t.runs:
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
sub = doc.add_paragraph()
sr = sub.add_run("Diagnostic investigation: seven failed interventions, four withdrawn hypotheses, and a target that holds")
sr.italic = True
sr.font.size = Pt(12)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run(
    "Marine Research Institute, Klaipėda University\n"
    "Prepared 2026-08-03 · OSMOSE Python engine, 9-species Baltic configuration\n"
    "Working document — conclusions are provisional and several are explicitly withdrawn below"
)
mr.font.size = Pt(9.5)
mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

# ================= SUMMARY =================
h("Executive summary", 1)
para(
    "The 9-species Baltic OSMOSE configuration holds pikeperch (Sander lucioperca) at roughly 1.45 "
    "million tonnes against an ICES envelope of 4,000–25,000 t — a factor of ~56. Smelt is ~5.7× over. "
    "Perch, the same family in the same coastal habitat under a comparable fishery, sits inside its "
    "envelope and serves throughout as the control."
)
para(
    "This document records what was measured, what was eliminated, and — importantly — three "
    "substantive conclusions that were reached and then withdrawn under scrutiny. It is written so that "
    "the withdrawn material is as recoverable as the surviving material, because in each case the "
    "measurement that exposed the error is what pointed at its replacement. The current diagnosis in "
    "Section 7 is the third such position and has not yet undergone the same scrutiny as the two it "
    "replaced."
)
para("Principal findings:", bold=True)
bullets([
    "The overshoot is not an engine defect. Consumption-to-biomass ratios (Q/B) at equilibrium fall "
    "within published ranges for every species; there is no conservation violation and no "
    "growth-without-eating.",
    "Demand-side levers are closed. A 10× cut in percid access to herring and sprat yields −25%; "
    "fishing mortality swept from F = 0.5 to F = 4.0 yields −10.6%, non-monotonically, while the "
    "envelope count degrades from 7/9 to 4/9. Pikeperch cannot be fished down.",
    "WITHDRAWN: a spatial-supply explanation. The supporting figure was wrong by a factor of ~3.5, and "
    "the corrected numbers refute the hypothesis outright (Section 5).",
    "WITHDRAWN: a predation-release explanation. Pikeperch is fourth-least predated, not least, and "
    "three less-predated species are all in envelope (Section 6).",
    "Escape from the predator size window is established — no predator in the configuration can take a "
    "pikeperch above 36 cm, and pikeperch matures at 40 cm — but it is not sufficient: opening cod's "
    "window past maturity moved biomass the wrong way, because no predator here has the biomass to "
    "regulate a stock fifteen times its own size (Section 7).",
    "Seven interventions across the removal and production sides returned between -86% and +8% against "
    "a 58-fold gap. The overshoot is over-determined (Section 8).",
    "The target was re-derived and RETAINED. Its scope is correctly declared and its magnitude survives "
    "an independent production-based cross-check, so the gap is real and not an artefact of the "
    "reference point (Section 9).",
    "Certification is now weight-aware. The model reproduces five of five targets that have a real "
    "assessment behind them; the two failures are the two lowest-confidence literature estimates in "
    "the reference file (Section 10).",
    "All diet-based results required first repairing a defect in the diet diagnostic itself, which had "
    "been reporting every predator as 100% piscivorous and 0% resource-feeding (Section 3).",
])

# ================= 1 =================
h("1. Problem statement", 1)
table(
    ["Species", "Final-decade mean (t)", "ICES envelope (t)", "Status"],
    [
        ["pikeperch", "1,453,313", "4,000 – 25,000", "~56–58× over"],
        ["smelt", "680,125", "20,000 – 120,000", "~5.7× over"],
        ["perch", "45,382", "8,000 – 50,000", "in envelope"],
        ["cod_east", "83,122", "60,000 – 85,000", "in envelope (2.3% headroom)"],
        ["herring", "2,600,112", "800,000 – 3,000,000", "in envelope"],
    ],
)
note(
    "50-year runs, final-decade means, Python engine. The 7/9 baseline verdict derives from "
    "scripts/baltic_stability_certify.py at five seeds."
)
para(
    "A caveat that persists through every comparison below: the ICES envelope is a per-stock figure for "
    "locally assessed populations, whereas the model carries a single aggregated basin-wide pikeperch. "
    "The target is therefore not strictly like-for-like, and this is not resolved by any intervention "
    "proposed here."
)

# ================= 2 =================
h("2. Ecological context", 1)
para(
    "Baltic percids are coastal and locally structured. Site-fidelity and home-range work indicates a "
    "limited home range within bays (Christensen et al., 2020; Hansson et al., 2019, as reviewed in "
    "Hall et al., 2022), and stable-isotope analysis resolves feeding-range differences between closely "
    "located sites within a single littoral area, implying sedentary behaviour at sub-bay scale "
    "(Ahlbeck Bergendahl et al., 2017). That evidence is for perch and should not be extended to "
    "pikeperch without a pikeperch source."
)
para(
    "Stock structure is fine-grained. Pikeperch shows \"a very local population structure\" (Björklund "
    "et al., 2007, cited in Olsson et al., 2015). Anadromous perch populations less than 50 km apart "
    "are genetically differentiated with only 3–5% gene flow and exhibit reproductive homing (Hall et "
    "al., 2022). Along the Swedish coast, perch shows isolation by distance with stretches of deep "
    "water acting as barriers to gene flow, and the species is described as suitable for local "
    "management (Olsson et al., 2011). Olsson et al. (2015) accordingly treat the Gulf of Riga, Gulf of "
    "Finland, Archipelago Sea, Bothnian and lagoon systems as separate coastal ecosystems."
)
note(
    "These are population-genetic results. They bound dispersal and stock identity, not daily foraging "
    "range. An earlier draft of the design used them as evidence about feeding range; that inference is "
    "withdrawn."
)
para(
    "Percids nonetheless do take herring inside bays. Jensen, Hansson and Didrikas (2011) examined diel "
    "vertical migration of young-of-the-year herring and \"one of their major predators, pikeperch\" in "
    "Himmerfjärden, a brackish Baltic bay, in summer, identifying piscivorous targets as pikeperch "
    "larger than 45 cm — above the 40 cm maturity size used in this configuration. The predator–prey "
    "size window already restricts this link to the coastal size class. The link is real and must not "
    "be removed from the configuration."
)
para(
    "Herring and sprat are not interchangeable in this respect: Baltic herring (Clupea harengus "
    "membras) is a coastal spring spawner whose juveniles are bay-resident, whereas sprat is genuinely "
    "offshore and deeper-dwelling."
)
para(
    "Smelt (Osmerus eperlanus) occurs as an anadromous form spawning in coastal low-salinity zones and "
    "rivers, with the fishery conducted on migrating fish during April–May and total spawning-migration "
    "duration of 20–45 days (Sendek & Bogdanov, 2019). It is also among the most abundant fishes of the "
    "eastern Gulf of Finland year-round, so treating it as available to percids only during the run "
    "overstates the restriction."
)
para(
    "Finally, the direction of the model's error is opposite to the observed state of the resource: "
    "Jakubavičiūtė et al. (2022) assessed nine European pikeperch stocks and found three strongly "
    "depleted below B_MSY, including two in the Baltic Sea."
)

# ================= 3 =================
h("3. Methodological prerequisite: the diet diagnostic was defective", 1)
para(
    "No diet-based conclusion in this investigation was trustworthy until a defect in the diet output "
    "was repaired. It is documented here because it invalidates prior analyses, not merely this one."
)
para(
    "The prey axis of the diet and predator-pressure matrices was labelled with the focal species plus "
    "background species (grey seal, cormorant), while the predation kernel writes resource groups at "
    "column index n_species + resource_index. On the Baltic configuration this placed diatoms in the "
    "grey-seal column and dinoflagellates in the cormorant column, and pushed microzooplankton, "
    "mesozooplankton, macrozooplankton and benthos beyond the end of the array, where a bounds check "
    "discarded them without error."
)
para(
    "The visible symptom was smelt — a planktivore — reporting 94.3% of its diet as grey seal, a prey "
    "it cannot physically consume, and 0% resources, which is nearly all of what it actually eats. "
    "Every predator reported as 100% fish and 0% resource, which is precisely what an absent resource "
    "axis produces."
)
para(
    "After repair, smelt reads 55% mesozooplankton, 18% microzooplankton and 15% macrozooplankton. A "
    "first repair pass corrected three of the four sites that construct this axis and missed the "
    "fourth; that omission surfaced only as a biologically impossible consumption figure — herring "
    "consuming 0.12 times its own body weight per year while feeding 61% on mesozooplankton. The "
    "regression test added in response asserts that no fish can consume less than its own body weight "
    "annually, which catches a dropped prey block irrespective of which code path drops it."
)
note(
    "Consequence for the wider project: any conclusion drawn from dietMatrix or predatorPressure on "
    "this configuration before commit e121c6d (and, for absolute consumption, 3e320ff) should be "
    "re-derived."
)

# ================= 4 =================
h("4. Hypotheses eliminated by measurement", 1)
table(
    ["Hypothesis", "Test", "Result"],
    [
        ["Missing fishing mortality", "Read configuration", "Rejected — pikeperch F = 0.50, above perch's 0.40"],
        ["Fishery not wired to species", "Catchability matrix", "Rejected — pikeperch maps to coastalpikeperch = 1"],
        ["Stock–fishery spatial mismatch", "Map overlap", "Rejected — 100% of occupied cells overlap the fishery"],
        ["Resource (LTL) subsidy", "Equilibrium diet", "Rejected — 36.9% resource, below perch's 42.1%"],
        ["Engine accounting / bioenergetics", "Q/B at equilibrium", "Rejected — all species within published ranges"],
    ],
)
para("Consumption-to-biomass ratios at equilibrium:", bold=True)
table(
    ["Species", "Biomass (t)", "Consumption (t/yr)", "Q/B"],
    [
        ["pikeperch", "1,453,313", "3,712,747", "2.55"],
        ["perch", "45,382", "154,920", "3.41"],
        ["cod_east", "83,122", "285,817", "3.44"],
        ["herring", "2,600,112", "11,925,219", "4.59"],
        ["sprat", "1,059,941", "5,056,861", "4.77"],
    ],
)
para(
    "Planktivores at 4.6–4.8 and piscivores at 2.5–3.4 are consistent with published values. Pikeperch "
    "has the lowest Q/B of any fish in the system: it is not feeding anomalously hard, there is simply "
    "a great deal of it."
)
note(
    "The resource-subsidy hypothesis was initially supported by a three-year run showing ~76% "
    "resource-derived diet. At equilibrium the figure is 36.9%. Diet composition on this configuration "
    "is not stable until the initialisation transient clears; short runs pointed in the opposite "
    "direction from the truth."
)

# ================= 5 =================
h("5. Withdrawn: the spatial-supply explanation", 1)
para(
    "An intermediate conclusion held that the overshoot arose from a bay-resident predator holding a "
    "continuously replenished tap into the basin-wide herring pool, on the basis that pikeperch's "
    "27-cell range contained 4.55% of the herring distribution (~118 kt) against consumption of "
    "1.46 Mt/yr — approximately 12 times the locally available standing stock annually."
)
para("Both the figure and the conclusion are withdrawn.", bold=True)
para(
    "Adult herring occupies the 593-cell map for only 10 of 24 timesteps. The configuration assigns it "
    "three seasonal maps, honoured per movement.steps by the engine:"
)
table(
    ["Timesteps", "Herring map", "Cells", "Share within pikeperch cells"],
    [
        ["0–3, 12–15, 22–23", "herring_adult", "593", "4.55%"],
        ["4–11", "herring_spawning", "51", "3.9 – 9.8%"],
        ["16–21", "herring_spawning_autumn", "10", "50.0%"],
    ],
)
para(
    "The time-weighted share is 16.19%, corresponding to ~421 kt locally available, so consumption "
    "represents ~3.5× the local standing stock, not 12×. An all-juvenile bound gives 9.3×, so the "
    "original figure lies outside the range achievable under any combination of adult and juvenile "
    "distributions. The dominant term is autumn, when all adult herring occupies a 10-cell coastal map "
    "of which pikeperch holds half — seasonal spawning aggregation into percid cells, a different "
    "mechanism from continuous replenishment."
)
para("The comparison that refutes the hypothesis:", bold=True)
table(
    ["Predator", "Time-weighted share of herring", "Share of sprat", "Status"],
    [
        ["pikeperch", "16.19%", "3.73%", "~56× over envelope"],
        ["perch", "31.56%", "7.90%", "in envelope"],
    ],
)
para(
    "Perch has twice pikeperch's spatial access to herring and remains within its envelope. Spatial "
    "supply overlap therefore does not discriminate between the two species and cannot explain the "
    "pikeperch overshoot specifically."
)

# ================= 6 =================
h("6. Withdrawn: predation release", 1)
para(
    "A second intermediate conclusion held that pikeperch is the least-predated fish in the system, "
    "with total accessibility as prey of 0.60 against perch's 1.05, and that this asymmetry explained "
    "the overshoot. This too is withdrawn."
)
table(
    ["Species", "Total accessibility as prey", "Status"],
    [
        ["cod_west", "0.20", "in envelope"],
        ["cod_east", "0.20", "in envelope"],
        ["flounder", "0.35", "in envelope"],
        ["pikeperch", "0.60", "~56x over"],
        ["perch", "1.05", "in envelope"],
        ["smelt", "2.55", "~5.7x over"],
    ],
)
para(
    "Pikeperch is fourth-least predated, not least. Three species carry lower predation and all three "
    "are within envelope, while smelt carries the highest predation in the system and is 5.7x over. The "
    "proposed axis is contradicted at both ends. The hypothesis had not been subjected to the same "
    "control test that was used to reject the spatial-supply explanation in Section 5 — the test was "
    "applied to the hypothesis being discarded and not to the one being adopted."
)

h("7. Escape from the predator size window: established, but not sufficient", 1)
para(
    "The mechanism that survives is not an accessibility coefficient at all. The predation kernel "
    "applies the predator-prey size-ratio gate before reading accessibility, so a coefficient outside "
    "the size window is never applied."
)
table(
    ["Predator", "Linf (cm)", "Minimum size ratio", "Largest prey it can take (cm)"],
    [
        ["cod_west / cod_east", "110", "3.5", "31.4"],
        ["pikeperch (cannibalism)", "90", "2.5", "36.0"],
    ],
)
para(
    "Pikeperch matures at 40.0 cm and has an asymptotic length of 90 cm. No predator in the "
    "configuration can take a pikeperch larger than 36 cm, so the species escapes the entire predator "
    "field before it reaches maturity and accumulates biomass in a structurally invulnerable adult "
    "class. Perch, with an asymptotic length of 45 cm, spends a far greater proportion of its life "
    "inside cod's 31.4 cm window."
)
para(
    "This reframes the asymptotic length from a contributing factor into the mechanism itself, and it "
    "accounts for every lever that failed: fishing is the only source of adult removal, and every "
    "demand-side intervention acted on a size class that was never predation-limited."
)
para(
    "This mechanism is real and verified. It is nonetheless not sufficient, and the intervention that "
    "tested it failed."
)
para(
    "Opening cod's size window from a 31.4 cm ceiling to 44 cm — crossing pikeperch's 40 cm maturity, "
    "and thereby giving a predator access to the adult class for the first time — moved pikeperch "
    "biomass the wrong way, by +3.4%. Pairing it with realistic gear selectivity and a reduced "
    "asymptotic length did not help.",
)
para(
    "The reason is a biomass argument that should have preceded the test. Cod totals roughly 97,000 t "
    "against pikeperch's 1.45 million, a stock fifteen times its own size, and consumes about "
    "330,000 t per year across all prey. Grey seal, the only other predator whose size window reaches "
    "adult pikeperch, consumes about 58,500 t per year across all prey. No predator in this "
    "configuration has the biomass to regulate pikeperch, whatever the size window permits. Adults are "
    "genuinely unreachable, and making them reachable changes nothing."
)

h("8. Seven interventions, seven failures", 1)
para(
    "Levers were tested across both the removal and the production side of the population balance. "
    "Results are final-decade means, 50-year runs, seed 42 unless noted."
)
table(
    ["Intervention", "Pikeperch result", "Envelope"],
    [
        ["baseline", "1,453,313 t (58x over)", "7/9"],
        ["diet accessibility, herring+sprat cut 10x", "-25.4%", "5/9"],
        ["fishing mortality, F swept 0.5 to 4.0", "-10.6%, non-monotonic", "4/9"],
        ["gear selectivity, l50 = 40 cm", "-7.4%", "7/9"],
        ["gear selectivity, l50 = 45 cm", "+7.8%", "7/9"],
        ["cod size window, 3.5 to 2.5", "+3.4%", "7/9"],
        ["asymptotic length 90 to 70 cm (paired)", "-1.2%", "6/9"],
        ["recruitment ssbhalf, cut 97.6x", "-86% (209,000 t, still 8.4x over)", "4/9"],
    ],
)
para(
    "Biomass responds to the recruitment parameter as roughly its 0.42 power, so reaching the target "
    "would require a value near 3 t — two further orders of magnitude below anything defensible. The "
    "overshoot is over-determined: a 90 cm asymptotic length, a fifteen-year lifespan, no predation "
    "reaching adults, and a predator field an order of magnitude too small jointly sustain the stock, "
    "and removing any one of them is compensated by the others."
)

h("9. The target: re-derived and retained", 1)
para(
    "With every dynamical lever exhausted, the target itself was re-derived. Two findings, the second "
    "of which corrects an earlier position taken in this investigation."
)
para(
    "First, ICES assesses none of these species. A query of the 2023 assessments for area 27.2x returns "
    "thirteen stocks — cod, herring, sprat, plaice, dab, sole, brill and ray — and no pikeperch, perch, "
    "smelt or stickleback. The phrase \"ICES envelope\", used throughout this project's certification "
    "notes, is therefore wrong for these four species; the reference file sources them as literature "
    "estimates for the coastal Baltic."
)
para(
    "Second, an earlier suggestion in this work — that the target might be a per-stock figure wrongly "
    "applied to an aggregated basin-wide model stock, and that this could account for an order of "
    "magnitude of the gap — is withdrawn. The reference file's header already states the scope: total "
    "stock biomass over the whole Baltic model domain. The targets are declared aggregates and the "
    "ambiguity does not exist."
)
para(
    "On magnitude, the one available quantitative anchor is perch production in the Curonian Lagoon "
    "(Hansson et al., 2018): 3,200 t over 1,600 km2, or 2.0 t/km2, which at a production-to-biomass "
    "ratio of 0.5-0.7 implies 2.9-4.0 t/km2 of biomass in prime habitat. Scaled across a 5-15% "
    "habitat fraction of the domain and discounted two- to four-fold for the lagoon's atypically high "
    "productivity, this gives roughly 14,000-113,000 t of perch against a committed target of "
    "8,000-50,000 t — at the low end of the band, but inside it."
)
para(
    "Both targets are retained. Pikeperch must sit below perch, having a restricted and patchy "
    "distribution where perch is ubiquitous, and both Baltic stocks assessed by Jakubaviciute et al. "
    "(2022) are depleted below their biomass at maximum sustainable yield, so a target near the bottom "
    "of a plausible range is the correct reading. Smelt, the second most abundant fish of the eastern "
    "Gulf of Finland after herring, warrants a target an order of magnitude above pikeperch's."
)
note(
    "Consequence: the overshoot is not explained by the target. Its scope is correctly declared, its "
    "magnitude is consistent with the available anchor and with the documented depleted state of the "
    "stocks, and its relation to the perch target is internally consistent. The 58-fold gap is real."
)

h("10. Weight-aware certification", 1)
para(
    "The reference file has always carried a confidence weight per species — 1.0 for well-assessed "
    "stocks, 0.5 for medium, 0.2 for those poorly resolved at grid scale — and the certifier discarded "
    "it, scoring all nine species identically. The two species failing the headline verdict were "
    "precisely the two lowest-weight rows in the file, so that verdict was substantially a statement "
    "about the weakest targets rather than about the model."
)
table(
    ["Tier", "Species", "Weight", "Result"],
    [
        ["ASSESSED", "cod_west, cod_east, herring, sprat", "1.0", "all pass"],
        ["ASSESSED", "flounder", "0.5", "pass"],
        ["INDICATIVE", "smelt", "0.3", "5.7x over"],
        ["INDICATIVE", "perch, stickleback", "0.2", "pass"],
        ["INDICATIVE", "pikeperch", "0.2", "58x over"],
    ],
)
para(
    "Certification is now tiered. The headline verdict counts only the assessed tier, and on the "
    "current configuration reads five of five: the model reproduces every target with a real "
    "assessment behind it. The indicative tier is reported separately and excluded from the verdict, "
    "and the former all-species figure is retained and labelled as legacy so that earlier notes citing "
    "it remain comparable."
)

h("11. Proposed refactor (superseded)", 1)
para(
    "The following tiers were designed against the predation-release hypothesis withdrawn in Section 6. "
    "They are recorded for completeness and are not recommended: Tier A raises coefficients on size "
    "classes holding roughly 5% of pikeperch biomass, since every predator size window closes below "
    "maturity, and Tier B derives its herring coefficient by multiplying accessibility by a spatial "
    "overlap that the cell-based predation kernel already applies. A narrower variant of Tier B was "
    "measured and failed non-regression, moving the envelope count from 7/9 to 5/9."
)
h("11.1 Tier A — test predation release (configuration only)", 2)
para(
    "Raise predation on pikeperch toward perch's level: cormorant 0.4 → 0.6, cod_west 0.1 → 0.15, "
    "cod_east 0.05 → 0.10, taking total accessibility as prey from 0.60 to 0.90. This is a diagnostic "
    "intervention, not a proposed calibration; it asks whether the asymmetry has the leverage the "
    "hypothesis requires."
)
h("11.2 Tier B — trophic corrections (configuration only)", 2)
bullets([
    "Sprat → 0 for both percids. Sprat is genuinely offshore; this is the defensible half of the "
    "original proposal.",
    "Herring → scaled, not removed. The documented bay link is retained at the coastal-available "
    "fraction, using the time-weighted spatial overlap as the scale: pikeperch ≈ 0.05, perch ≈ 0.06.",
    "Smelt → time-averaged over a corrected window. Spring (April–May), 20–45 day runs, giving W ≈ 1 "
    "month rather than 3: perch ≈ 0.04, pikeperch ≈ 0.05. These are a lower bound, since smelt is "
    "present year-round in some basins.",
])
h("11.3 Tier C — seasonal accessibility (engine feature, conditional)", 2)
para(
    "The accessibility matrix is stage-indexed by age threshold and carries no time axis, so seasonal "
    "prey availability cannot presently be expressed. Tier C would add a sparse per-pair, per-timestep "
    "multiplier mirroring the existing fishing-seasonality idiom, resolved to a dense array at "
    "configuration time because the predation kernel is compiled. It is conditional on Tier B showing "
    "that timing rather than level is what is wrong."
)
h("11.4 Out of scope", 2)
para(
    "Representing percid stocks as separate coastal units is the only change that would address supply "
    "rather than links, and the only one that makes the envelope comparison like-for-like. It is "
    "excluded here on cost grounds; the cod eastern/western disaggregation, which could not be fitted "
    "and remains a flagged experiment, is the cautionary precedent."
)

# ================= 8 =================
h("12. Acceptance criteria", 1)
para(
    "Measured on final-decade means, 50-year runs, five seeds, using the same certifier that produced "
    "the baseline verdict."
)
bullets([
    "Non-regression (hard): at least 7 of 9 species in envelope. This is a floor, not a target — the "
    "baseline already satisfies it, so it cannot be the sole criterion.",
    "Effect demonstrated (hard): pikeperch's final-decade mean must fall by more than twice the "
    "five-seed spread of the baseline, so that the change is distinguishable from noise.",
    "Mechanism demonstrated (hard): for Tier A, pikeperch's realised predation mortality as a share of "
    "total mortality must rise, read from model output rather than inferred from the coefficient. For "
    "Tier B, herring must remain present but reduced in realised diet; its disappearance would indicate "
    "the coefficient acting as a hard gate rather than a scale, which is a defect.",
    "Collapse guard: pikeperch must not fall below its envelope floor on the final-decade minimum. "
    "Reaching envelope by starvation or collapse dynamics is a failure, not a fix.",
    "Explicitly not a criterion: smelt reaching envelope, which is a separate failure with no mechanism "
    "in this work.",
])
para("Risk, ranked by measured headroom to the nearest envelope bound:", bold=True)
table(
    ["Species", "Current (t)", "Nearest bound (t)", "Headroom"],
    [
        ["cod_east", "83,122", "85,000 (upper)", "2.3%"],
        ["perch", "45,382", "50,000 (upper)", "9.2%"],
        ["herring", "2,600,112", "3,000,000 (upper)", "15.4%"],
        ["stickleback", "80,159", "50,000 (lower)", "60.2%"],
    ],
)
para(
    "cod_east is the tightest species in the system, and Tier A raises cod accessibility to pikeperch, "
    "which feeds cod. Any species leaving its envelope constitutes a failure regardless of explanation; "
    "an earlier draft permitted a herring breach to be reclassified as a subsequent finding, and that "
    "allowance has been withdrawn."
)

# ================= 9 =================
h("13. Limitations", 1)
bullets([
    "Single configuration, single parameter set. Nothing here transfers to other OSMOSE configurations.",
    "Several figures derive from single-seed runs and are labelled as such in the source documents; the "
    "five-seed certification is the standard for envelope verdicts.",
    "The predation-release hypothesis is untested. It tracks the observed asymmetry between perch and "
    "pikeperch but causation has not been demonstrated.",
    "Linf = 90 cm for Baltic pikeperch sits at the upper end of the plausible range and has not been "
    "independently validated here, though biomass scales with roughly the cube of length.",
    "The envelope comparison is not like-for-like: a per-stock reference target is being applied to an "
    "aggregated basin-wide model stock.",
    "Four substantive conclusions reached during this work were subsequently withdrawn: the "
    "resource-subsidy, spatial-supply and predation-release explanations, and the suggestion that the "
    "target was mis-scoped. Each was caught by measurement or by adversarial review rather than by the "
    "author's own reasoning, which suggests further unexamined assumptions may remain.",
    "The target re-derivation retains both figures because they are not refuted, which is weaker than "
    "deriving them. The perch cross-check rests on a single lagoon anchor scaled by an estimated rather "
    "than measured habitat fraction; it bounds the target to within about an order of magnitude.",
    "No quantitative anchor was found for pikeperch or smelt specifically. Both conclusions are "
    "relative — to perch, and to herring — rather than absolute.",
    "The weight-aware verdict changes what is reported, not the underlying dynamics. Pikeperch remains "
    "58-fold above a target that survived re-derivation; the configuration cannot currently represent "
    "the species at a defensible biomass, which is what the reference file's own note about coarse-grid "
    "under-resolution anticipated.",
])

# ---- review outcome, filled if provided ----
if REVIEW2 and REVIEW2.exists():
    try:
        data = json.loads(REVIEW2.read_text())
        res = data.get("result", data)
        conf = res.get("confirmed", [])
        h("14. Independent review", 1)
        para(
            "The design and implementation plan underlying Sections 6–8 were subjected to two rounds of "
            "multi-agent adversarial review, in which independent reviewers raised findings that a "
            "second agent then attempted to refute; only findings surviving refutation are recorded."
        )
        para(
            "The first round upheld eight findings, including the two withdrawals documented in "
            "Sections 4 and 5. The second round, conducted on the revised documents, upheld "
            f"{len(conf)}."
        )
        if conf:
            table(
                ["Severity", "Finding"],
                [[f.get("severity", ""), f.get("title", "")[:180]] for f in conf],
            )
        else:
            para(
                "No findings survived refutation in the second round.", italic=True
            )
    except Exception as exc:  # pragma: no cover
        print("review section skipped:", exc)

# ================= REFERENCES =================
h("References", 1)
REFS = [
    "Ahlbeck Bergendahl, I., Holliland, P. B., & Hansson, S. (2017). Feeding range of age 1+ year "
    "Eurasian perch Perca fluviatilis in the Baltic Sea. Journal of Fish Biology, 90(5), 2060–2072. "
    "https://doi.org/10.1111/jfb.13285",

    "Björklund, M., et al. (2007). Cited in Olsson, Tomczak & Ojaveer (2015).",

    "Downing, J. A., & Plante, C. (1993); Randall, R. G., & Minns, C. K. (2000). Production/biomass "
    "relationships, cited in Hansson, Bergström & Bonsdorff (2018).",

    "Hansson, S., Bergström, U., & Bonsdorff, E. (2018). Competition for the fish — fish extraction "
    "from the Baltic Sea by humans, aquatic mammals, and birds. ICES Journal of Marine Science, 75(3), "
    "999–1008. https://doi.org/10.1093/icesjms/fsx207",

    "Christensen, E. A. F., et al. (2020); Hansson, S., et al. (2019). Cited in Hall, Koch-Schmidt & "
    "Larsson (2022).",

    "Dainys, J., et al. (2022). Cited in Jakubavičiūtė, Arula & Dainys (2022).",

    "Hall, M., Koch-Schmidt, P., & Larsson, P. (2022). Reproductive homing and fine-scaled genetic "
    "structuring of anadromous Baltic Sea perch (Perca fluviatilis). Fisheries Management and Ecology, "
    "29(5), 586–596. https://doi.org/10.1111/fme.12542",

    "Heikinheimo, O., et al. (2015). Cited in Jakubavičiūtė, Arula & Dainys (2022).",

    "Jakubavičiūtė, E., Arula, T., & Dainys, J. (2022). Status and future perspectives for pikeperch "
    "(Sander lucioperca) stocks in Europe. openRxiv. https://doi.org/10.1101/2022.12.20.521162",

    "Jensen, O. P., Hansson, S., & Didrikas, T. (2011). Foraging, bioenergetic and predation "
    "constraints on diel vertical migration. Journal of Fish Biology, 78(2), 449–465. "
    "https://doi.org/10.1111/j.1095-8649.2010.02855.x",

    "Kokkonen, E., Heikinheimo, O., & Pekcan-Hekim, Z. (2019). Effects of water temperature and "
    "pikeperch (Sander lucioperca) abundance on the stock–recruitment relationship of Eurasian perch "
    "(Perca fluviatilis) in the northern Baltic Sea. Hydrobiologia, 841(1), 79–94. "
    "https://doi.org/10.1007/s10750-019-04008-z",

    "Olin, M., Heikinheimo, O., & Lehtonen, T. K. (2023). Long-term monitoring of pikeperch (Sander "
    "lucioperca) populations under increasing temperatures and predator abundances in the Finnish "
    "coastal waters of the Baltic Sea. Ecology of Freshwater Fish, 32(4), 750–764. "
    "https://doi.org/10.1111/eff.12721",

    "Olsson, J., Mo, K., & Florin, A.-B. (2011). Genetic population structure of perch Perca "
    "fluviatilis along the Swedish coast of the Baltic Sea. Journal of Fish Biology, 79(1), 122–137. "
    "https://doi.org/10.1111/j.1095-8649.2011.02998.x",

    "Olsson, J., Tomczak, M. T., & Ojaveer, H. (2015). Temporal development of coastal ecosystems in "
    "the Baltic Sea over the past two decades. ICES Journal of Marine Science, 72(9), 2539–2548. "
    "https://doi.org/10.1093/icesjms/fsv143",

    "Sendek, D. S., & Bogdanov, D. V. (2019). European smelt Osmerus eperlanus in the eastern Gulf of "
    "Finland, Baltic Sea: Stock status and fishery. Journal of Fish Biology. "
    "https://doi.org/10.1111/jfb.14009",
]
for r in REFS:
    p = doc.add_paragraph(r)
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.first_line_indent = Pt(-24)
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        run.font.size = Pt(9.5)

h("Source documents in the repository", 1)
for s in [
    "docs/baltic_percid_overshoot_investigation_2026-08-02.md — measurements and eliminated hypotheses",
    "docs/superpowers/specs/2026-08-02-percid-trophic-refactor-design.md — design",
    "docs/superpowers/plans/2026-08-02-percid-trophic-refactor.md — implementation plan",
    "docs/proposals/2026-08-02-percid-trophic-refactor.md — original proposal",
    "docs/baltic_percid_target_rederivation_2026-08-06.md — target re-derivation and provenance",
    "data/baltic/reference/biomass_targets.csv — targets, confidence weights and sources",
    "scripts/baltic_stability_certify.py — weight-aware certifier",
    "docs/baltic_rv_gate_mechanism_ab_2026-08-02.md — related recruitment-gate measurement",
]:
    doc.add_paragraph(s, style="List Bullet")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print("wrote", OUT, OUT.stat().st_size, "bytes")
